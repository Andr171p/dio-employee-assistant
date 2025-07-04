from docx import Document as DOCXDocument
import base64

from docx2md import DocxFile, DocxMedia, Converter

from src.employee_assistant.base import LoadFileError
from src.employee_assistant.schemas import MarkdownDocument, ImageMetadata
from src.employee_assistant.utils import get_file_extension, get_file_size_by_base64

file_path = r"C:\Users\andre\IdeaProjects\DIORag\knowledge_base\Инструкции\ИНСТРУКЦИЯ_по_заполнению_в_1С_УФФ_документа_Задание_сотруднику_29.docx"


def extract_images(file_path: str) -> ...:
    document = DOCXDocument(file_path)
    # images: list[str] = []
    images: dict[str, str] = {}
    id = 0
    for rel_value in document.part.rels.values():
        if "image" in rel_value.target_ref:
            image_data = rel_value.target_part.blob
            image_base64 = base64.b64encode(image_data).decode("utf-8")
            # images.append(image_base64)
            id += 1
            images[f"image{id}"] = image_base64
    return images


imgs = extract_images(file_path)

print(imgs.keys())


class DOCX2MarkdownLoader:
    def __init__(self, file_path: str, **kwargs) -> None:
        self.file_path = file_path
        self.use_md_table: bool = kwargs.get("use_md_table", False)

    def load(self) -> MarkdownDocument:
        try:
            images = self._extract_images()
            docx = DocxFile(str(file_path))
            media = DocxMedia(docx)
            converter = Converter(docx.document(), media, self.use_md_table)
            md_text = converter.convert()
            return MarkdownDocument(
                source=file_path.split("/")[-1],
                content=md_text,
                additional_files=images
            )
        except Exception as e:
            raise LoadFileError(f"Error while loading document: {e}") from e

    def _extract_images(self) -> list[ImageMetadata]:
        document = DOCXDocument(self.file_path)
        images: list[ImageMetadata] = []
        id = 0
        for _, rel in document.part.rels.items():
            if "image" not in rel.target_ref:
                continue
            image_src = rel.target_ref.split("/")[-1]
            image_data = rel.target_part.blob
            image_base64 = base64.b64encode(image_data).decode("utf-8")
            id += 1
            image = ImageMetadata(
                file_path=image_src,
                format=get_file_extension(image_src),
                size_mb=get_file_size_by_base64(image_base64),
                file_base64=image_base64,
                image_id=f"image{id}"
            )
            images.append(image)
        return images
