from typing import Optional
from collections.abc import Iterator

import os
import base64
from pathlib import Path

from bs4 import BeautifulSoup

from langchain_core.documents import Document
from langchain_core.document_loaders import BaseLoader
from langchain_text_splitters import TextSplitter, RecursiveCharacterTextSplitter

from .base import ImageMetadata, get_size_by_base64


class Docx2MdLoader(BaseLoader):
    """Loader for DOCX files that converts them to Markdown format.

    Args:
        docx_path: Path to the DOCX file
        use_md_table: Whether to use Markdown tables (default: False)
        include_images: Whether to include images as base64 (default: False)
    """
    def __init__(
            self,
            docx_path: Path | str,
            use_md_table: bool = False,
            include_images: bool = False
    ) -> None:
        self.docx_path = str(docx_path)
        self.use_md_table = use_md_table
        self.include_images = include_images

    @property
    def file_name(self) -> str:
        return Path(self.docx_path).name

    @property
    def file_type(self) -> str:
        return self.file_name.split(".")[-1]

    def lazy_load(self) -> Iterator[Document]:
        try:
            from docx2md import DocxFile, DocxMedia, Converter
        except ImportError as ex:
            raise ImportError(
                "Could not import docx2md python package. "
                "Please install it with `pip install docx2md`."
            ) from ex
        images = self._load_images() if self.include_images else []
        try:
            docx = DocxFile(self.docx_path)
            media = DocxMedia(docx)
            converter = Converter(docx.document(), media, use_md_table=self.use_md_table)
            md_text = converter.convert()
            metadata = {
                "source": self.docx_path,
                "file_path": self.docx_path,
                "file_name": self.file_name,
                "file_type": self.file_type,
                "images": images
            }
            yield Document(page_content=md_text, metadata=metadata)
        except Exception as ex:
            raise RuntimeError(f"Error while loading documents: {ex}") from ex

    def load_and_split(self, text_splitter: Optional[TextSplitter] = None) -> list[Document]:
        text_splitter = text_splitter or RecursiveCharacterTextSplitter()
        documents = self.load()
        if not self.include_images:
            return text_splitter.split_documents(documents)
        chunks: list[Document] = []
        for document in documents:
            texts = text_splitter.split_text(document.page_content)
            images = document.metadata["images"]
            for text in texts:
                srcs = self.__parse_image_src_from_md_text(text)
                if not srcs:
                    continue
                in_text_images = [image for image in images if image["src"] in srcs]
                metadata = document.metadata
                metadata["images"] = in_text_images
                chunks.append(Document(page_content=text, metadata=metadata))
        return chunks

    def _load_images(self) -> list[ImageMetadata]:
        try:
            from docx import Document as DOCXDocument
        except ImportError as ex:
            raise ImportError(
                "Could not import docx python package. "
                "Please install it with `pip install python-docx`."
            ) from ex
        if not os.path.exists(self.docx_path):
            raise ValueError(f"Path {self.docx_path} does not exists")
        docx_document = DOCXDocument(self.docx_path)
        images: list[ImageMetadata] = []
        for _, rel in docx_document.part.rels.items():
            if "image" not in rel.target_ref:
                continue
            src = rel.target_ref.split("/")[-1]
            data = rel.target_part.blob
            str_base64 = base64.b64encode(data).decode("utf-8")
            size_mb = get_size_by_base64(str_base64)
            extension = src.split(".")[-1].lower()
            images.append(ImageMetadata(
                src=src,
                source=self.docx_path,
                format=extension,
                size_mb=size_mb,
                str_base64=str_base64
            ))
        return images

    @staticmethod
    def __parse_image_src_from_md_text(md_text: str) -> list[str]:
        soup = BeautifulSoup(md_text, "html.parser")
        tags = soup.find_all("img")
        return [img["src"].split("/")[-1] for img in tags if "src" in img.attrs]
